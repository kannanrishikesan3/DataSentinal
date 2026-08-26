"""DOCX parser (python-docx) — paragraph- and table-cell-addressable
(reported via `line_number`, since DOCX has no native page concept without
a full layout engine)."""

from __future__ import annotations

from collections.abc import Iterator
from pathlib import Path

from datasentinel_agent.config.scan_config import load_scan_config
from datasentinel_agent.parsers.archive_guard import check_zip_safety
from datasentinel_agent.parsers.base import DocumentParser, ExtractedUnit, ParserError


class DocxParser(DocumentParser):
    def can_parse(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def extract_units(self, file_path: Path) -> Iterator[ExtractedUnit]:
        limits = load_scan_config().archive_limits
        check_zip_safety(
            file_path,
            max_members=limits.max_members,
            max_uncompressed_bytes=limits.max_uncompressed_bytes,
            max_ratio=limits.max_ratio,
        )

        try:
            import docx
        except ImportError as exc:
            raise ParserError("python-docx is not installed") from exc

        try:
            document = docx.Document(str(file_path))
        except Exception as exc:
            raise ParserError(f"Failed to open DOCX {file_path}: {exc}") from exc

        line_number = 0
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                line_number += 1
                yield ExtractedUnit(text=paragraph.text, line_number=line_number)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        line_number += 1
                        yield ExtractedUnit(text=cell.text, line_number=line_number)
